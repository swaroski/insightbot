import os
import uuid
from typing import Dict, Any, Optional
import PyPDF2
import docx
from loguru import logger
from app.services.vector_store import vector_store
from app.models.database import Document, SessionLocal
from app.config import settings


class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._process_pdf,
            'text/plain': self._process_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
        }
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_text(self, file_path: str) -> str:
        """Extract text from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise
    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def process_document(
        self, 
        file_path: str, 
        filename: str, 
        content_type: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """Process a document and add it to the vector store"""
        try:
            # Check if content type is supported
            if content_type not in self.supported_formats:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Extract text content
            processor = self.supported_formats[content_type]
            content = processor(file_path)
            
            if not content.strip():
                raise ValueError("No text content found in document")
            
            # Create database record
            document_id = str(uuid.uuid4())
            file_size = os.path.getsize(file_path)
            
            # Add to vector store
            vector_document_id = vector_store.add_document(
                content=content,
                filename=filename,
                metadata=metadata or {}
            )
            
            # Save to database
            db = SessionLocal()
            try:
                db_document = Document(
                    id=document_id,
                    filename=filename,
                    content_type=content_type,
                    file_size=file_size,
                    processed=True,
                    embedding_count=vector_store.get_chunk_count()
                )
                db.add(db_document)
                db.commit()
                
                logger.info(f"Successfully processed document {filename}")
                return document_id
                
            finally:
                db.close()
                
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise
    
    def get_supported_formats(self) -> list:
        """Get list of supported content types"""
        return list(self.supported_formats.keys())
    
    def is_supported_format(self, content_type: str) -> bool:
        """Check if content type is supported"""
        return content_type in self.supported_formats


# Global document processor instance
document_processor = DocumentProcessor()